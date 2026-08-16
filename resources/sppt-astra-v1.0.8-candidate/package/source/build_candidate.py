#!/usr/bin/env python3
"""Rebuild the SPPT/ASTRA v1.0.8 candidate paper editions.

This script is document engineering only. It does not modify the ASTRA GitHub
repository, run repository CI, or create a release identity.
"""

from __future__ import annotations

import argparse
import base64
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree

import matplotlib.font_manager as font_manager
import pikepdf
import pypandoc
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from playwright.sync_api import sync_playwright
from pypdf import PdfReader, PdfWriter

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "source"
FIGURES = ROOT / "figures"
VERIFY = ROOT / "verification"
MD = SOURCE / "ASTRA_SPPT_v1.0.8_Endogenous_Visibility_Candidate.md"
BUILD_EPOCH = "2026-08-16T00:00:00Z"
PDF_DATE = "D:20260816000000+00'00'"
WORK: Path


def _required_executable(environment_name: str, fallback: str | None = None) -> Path:
    configured = os.environ.get(environment_name)
    candidate = configured or (shutil.which(fallback) if fallback else None)
    if not candidate:
        raise RuntimeError(f"{environment_name} must name an executable")
    path = Path(candidate).resolve()
    if not path.is_file():
        raise RuntimeError(f"{environment_name} is not a file: {path}")
    return path


def pandoc_executable() -> Path:
    configured = os.environ.get("ASTRA_PANDOC")
    path = Path(configured).resolve() if configured else Path(pypandoc.get_pandoc_path()).resolve()
    if not path.is_file():
        raise RuntimeError(f"Pandoc executable is unavailable: {path}")
    return path


def tex_engine() -> Path:
    return _required_executable("ASTRA_TEX_ENGINE", "tectonic")


def fontconfig_file() -> Path:
    """Create an isolated Fontconfig binding for the locked Matplotlib fonts."""
    font_directory = Path(
        font_manager.findfont(
            font_manager.FontProperties(family="DejaVu Serif"),
            fallback_to_default=False,
        )
    ).parent
    config_directory = WORK / "fontconfig"
    config_directory.mkdir(parents=True, exist_ok=True)
    cache_directory = config_directory / "cache"
    cache_directory.mkdir(exist_ok=True)
    path = config_directory / "fonts.conf"
    path.write_text(
        '<?xml version="1.0"?>\n'
        '<!DOCTYPE fontconfig SYSTEM "urn:fontconfig:fonts.dtd">\n'
        "<fontconfig>\n"
        f"  <dir>{font_directory.as_posix()}</dir>\n"
        f"  <cachedir>{cache_directory.as_posix()}</cachedir>\n"
        "</fontconfig>\n",
        encoding="utf-8",
        newline="\n",
    )
    return path


def run(*args: str) -> None:
    environment = os.environ.copy()
    config = fontconfig_file()
    environment.update(
        {
            "FONTCONFIG_FILE": str(config),
            "FONTCONFIG_PATH": str(config.parent),
            "PYTHONHASHSEED": "0",
            "SOURCE_DATE_EPOCH": "1786838400",
            "TZ": "UTC",
        }
    )
    subprocess.run(args, cwd=ROOT, env=environment, check=True)


def run_pandoc(*args: str) -> None:
    run(str(pandoc_executable()), *args)


def normalize_docx_archive(path: Path) -> None:
    """Remove unverified renderer counts and fix package-member timestamps."""
    with zipfile.ZipFile(path) as archive:
        members = {
            info.filename: (info, archive.read(info.filename)) for info in archive.infolist()
        }

    app_name = "docProps/app.xml"
    if app_name in members:
        info, payload = members[app_name]
        root = ElementTree.fromstring(payload)
        for local_name in (
            "Characters",
            "CharactersWithSpaces",
            "Lines",
            "Pages",
            "Paragraphs",
            "TotalTime",
            "Words",
        ):
            for element in list(root):
                if element.tag.rsplit("}", 1)[-1] == local_name:
                    root.remove(element)
        for element in root:
            if element.tag.rsplit("}", 1)[-1] == "Application":
                element.text = "Pandoc 3.6.1 and python-docx 1.2.0"
        members[app_name] = (
            info,
            ElementTree.tostring(root, encoding="utf-8", xml_declaration=True),
        )

    temporary = path.with_suffix(".normalized.docx")
    with zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as archive:
        for name in sorted(members):
            original, payload = members[name]
            info = zipfile.ZipInfo(name, date_time=(2026, 8, 16, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.create_system = original.create_system
            archive.writestr(info, payload)
    temporary.replace(path)


def prepare_body_text(text: str) -> str:
    """Suppress the duplicate Pandoc title block before prepending text covers."""
    lines = text.splitlines()
    out: list[str] = []
    in_yaml = False
    yaml_done = False
    for line in lines:
        if not yaml_done and line.strip() == "---":
            if not in_yaml:
                in_yaml = True
                out.append(line)
                continue
            in_yaml = False
            yaml_done = True
            out.append(line)
            continue
        if in_yaml and any(
            line.startswith(key) for key in ("title:", "subtitle:", "author:", "date:")
        ):
            continue
        out.append(line)
    # Both PDF and DOCX builders prepend native, selectable text-only covers.
    return "\n".join(out) + "\n"


def _docx_static_toc(text: str) -> str:
    """Build a deterministic chapter-level TOC for headless DOCX rendering.

    Pandoc's live TOC field is not reliably materialized across DOCX renderers.
    A static hierarchy is therefore inserted for the DOCX only; the PDF keeps
    its native linked TOC and outline.
    """
    headings: list[str] = []
    in_yaml = False
    yaml_done = False
    for line in text.splitlines():
        if not yaml_done and line.strip() == "---":
            if not in_yaml:
                in_yaml = True
            else:
                in_yaml = False
                yaml_done = True
            continue
        if in_yaml:
            continue
        match = re.match(r"^# (.+?)(?: \{[^}]*\})?$", line.strip())
        if not match:
            continue
        title = match.group(1).strip()
        if title == "Table of Contents":
            continue
        headings.append(title)

    front: list[str] = []
    parts: list[tuple[str, list[str]]] = []
    appendices: list[str] = []
    current_part: tuple[str, list[str]] | None = None
    for title in headings:
        if title.startswith("Part "):
            current_part = (title, [])
            parts.append(current_part)
        elif title.startswith("Appendix "):
            appendices.append(title)
            current_part = None
        elif current_part is not None and re.match(r"^\d+\.", title):
            current_part[1].append(title)
        elif not re.match(r"^\d+\.", title):
            front.append(title)
        elif parts:
            parts[-1][1].append(title)
        else:
            front.append(title)

    lines = ["# Table of Contents {-}", ""]
    for title in front:
        lines.append(f"- {title}")
    for part, chapters in parts:
        lines.append(f"- **{part}**")
        for chapter in chapters:
            # Escape the numeric dot so Pandoc does not turn a nested bullet
            # into a second ordered-list block with a dangling dash.
            chapter_display = re.sub(r"^(\d+)\.", r"\1\\.", chapter)
            lines.append(f"  - {chapter_display}")
    if appendices:
        lines.append("- **Appendices**")
        for title in appendices:
            lines.append(f"  - {title}")
    lines.append("")
    return "\n".join(lines)


def prepare_docx_body_text(text: str) -> str:
    """Prepare body Markdown with a materialized, static DOCX contents list."""
    body = prepare_body_text(text)
    # Disable Pandoc's live TOC field for the DOCX; it renders as an empty
    # heading in independent DOCX renderers.
    body = re.sub(r"(?m)^toc:[ \t]*true[ \t]*$", "toc: false", body)
    body = re.sub(r"(?m)^toc-depth:[ \t]*\d+[ \t]*$", "", body)
    marker = "# Abstract {-}"
    if marker not in body:
        raise RuntimeError("Abstract marker not found while inserting DOCX TOC")
    return body.replace(marker, _docx_static_toc(body) + "\n" + marker, 1)


def build_pdf() -> None:
    body_md = WORK / "body_without_cover.md"
    body_md.write_text(prepare_body_text(MD.read_text(encoding="utf-8")), encoding="utf-8")
    body_pdf = WORK / "ASTRA_SPPT_v1.0.8_body.pdf"
    run_pandoc(
        str(body_md),
        "--from",
        "markdown+raw_tex",
        "--citeproc",
        "--bibliography",
        str(SOURCE / "references.bib"),
        f"--pdf-engine={tex_engine()}",
        f"--resource-path={os.pathsep.join((str(SOURCE), str(FIGURES)))}",
        "-o",
        str(body_pdf),
    )

    # Build a selectable, text-only cover. No scientific diagram is used.
    cover_pdf = WORK / "cover_only.pdf"
    cover_tex = WORK / "cover_only.tex"
    cover_tex.write_text(
        r"""\documentclass[10pt]{article}
\usepackage[paperwidth=8.5in,paperheight=11in,margin=0in]{geometry}
\usepackage{xcolor,fontspec}
\setmainfont{DejaVu Sans}
\definecolor{Navy}{HTML}{071B2A}
\definecolor{Gold}{HTML}{B99822}
\definecolor{Slate}{HTML}{3B4B56}
\pagestyle{empty}
\begin{document}
\begin{center}
\vspace*{1.25in}
{\color{Gold}\sffamily\bfseries\Large SPPT / ASTRA}\par
\vspace{0.65in}
{\color{Navy}\sffamily\bfseries\fontsize{42}{46}\selectfont v1.0.8}\par
\vspace{0.16in}
{\color{Navy}\sffamily\bfseries\fontsize{23}{27}\selectfont ENDOGENOUS VISIBILITY}\par
\vspace{0.34in}
{\color{Gold}\rule{5.6in}{1.2pt}}\par
\vspace{0.48in}
{\color{Navy}\sffamily\bfseries\fontsize{23}{28}\selectfont SOURCE-COUPLED TRANSDUCERS\par AND TRANSFORMED ARCHIVES}\par
\vspace{0.46in}
{\color{Slate}\sffamily\fontsize{11.5}{17}\selectfont Source-coupled transducers · self-detuning media\par cross-channel rescue · catastrophic archives · global certificates}\par
\vfill
{\color{Slate}\sffamily\fontsize{10}{15}\selectfont Candidate successor draft\par Not peer reviewed · no empirical planetary validation\par Repository-visible · no tag, release, Pages, DOI, or Zenodo action}\par
\vspace{0.36in}
\begin{minipage}{7.2in}
{\color{Slate}\sffamily 16 August 2026}\hfill{\color{Gold}\sffamily\bfseries AD ASTRA PER ASPERA}\par
\end{minipage}
\vspace*{0.65in}
\end{center}
\end{document}
""",
        encoding="utf-8",
    )
    engine = tex_engine()
    if engine.stem.casefold() == "tectonic":
        run(str(engine), "--outdir", str(WORK), str(cover_tex))
    else:
        run(
            str(engine),
            "-interaction=nonstopmode",
            "-halt-on-error",
            "-output-directory",
            str(WORK),
            str(cover_tex),
        )

    writer = PdfWriter()
    cover_reader = PdfReader(str(cover_pdf))
    body_reader = PdfReader(str(body_pdf))
    for page in cover_reader.pages:
        writer.add_page(page)
    for page in body_reader.pages:
        writer.add_page(page)

    # Preserve the manuscript's navigation hierarchy after the cover is
    # prepended. pypdf does not carry outlines through page copying unless
    # they are reconstructed explicitly.
    writer.add_outline_item("Cover", 0)

    def copy_outline(items, parent=None):
        last_item = None
        for item in items:
            if isinstance(item, list):
                copy_outline(item, last_item if last_item is not None else parent)
                continue
            page_number = body_reader.get_destination_page_number(item)
            if page_number is None or page_number < 0:
                continue
            last_item = writer.add_outline_item(
                getattr(item, "title", str(item)),
                page_number + 1,
                parent=parent,
            )

    copy_outline(body_reader.outline)
    writer.add_metadata(
        {
            "/Title": "SPPT / ASTRA v1.0.8 Candidate - Endogenous Visibility",
            "/Author": "Jacko T.",
            "/Subject": "Stateful edges, active supports, nonreciprocity, visibility, and ASTRA integration candidate",
            "/Keywords": "SPPT, ASTRA, stateful edges, nonreciprocity, active support, planetary inference",
        }
    )
    merged_pdf = WORK / "candidate-merged.pdf"
    with merged_pdf.open("wb") as f:
        writer.write(f)
    _sanitize_pdf(
        merged_pdf,
        ROOT / "ASTRA_SPPT_v1.0.8_Endogenous_Visibility_Candidate.pdf",
        "SPPT / ASTRA v1.0.8 Candidate: Endogenous Visibility",
    )


def style_docx(path_in: Path, path_out: Path) -> None:
    doc = Document(str(path_in))
    cp = doc.core_properties
    cp.title = "SPPT / ASTRA v1.0.8 Candidate - Endogenous Visibility"
    cp.subject = "Stateful edges, active supports, nonreciprocal effective interactions, sector-complete instruments, and operator-aware visibility"
    cp.author = "Jacko T."
    cp.keywords = "SPPT, ASTRA, stateful edges, nonreciprocity, active support, planetary inference"
    cp.comments = (
        "Candidate successor manuscript; not peer reviewed; no empirical planetary validation."
    )

    navy = RGBColor(0x07, 0x1B, 0x2A)
    slate = RGBColor(0x3B, 0x4B, 0x56)

    for sec in doc.sections:
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.top_margin, sec.bottom_margin = Inches(0.82), Inches(0.78)
        sec.left_margin = sec.right_margin = Inches(0.88)
        sec.header_distance = sec.footer_distance = Inches(0.32)
        sec.different_first_page_header_footer = True

    styles = doc.styles

    def set_font(name: str, font: str, size: float, bold: bool | None, color: RGBColor | None):
        st = styles[name]
        st.font.name = font
        st._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        st.font.size = Pt(size)
        if bold is not None:
            st.font.bold = bold
        if color is not None:
            st.font.color.rgb = color
        return st

    normal = set_font("Normal", "DejaVu Serif", 11.0, False, RGBColor(0x1C, 0x24, 0x2A))
    normal.paragraph_format.space_after = Pt(6.0)
    normal.paragraph_format.line_spacing = 1.24
    if "Body Text" in styles:
        body = set_font("Body Text", "DejaVu Serif", 11.0, False, RGBColor(0x1C, 0x24, 0x2A))
        body.paragraph_format.space_after = Pt(6.0)
        body.paragraph_format.line_spacing = 1.24
    set_font("Title", "DejaVu Sans", 25, True, navy).paragraph_format.space_after = Pt(4)
    set_font("Subtitle", "DejaVu Sans", 12.0, False, slate)
    set_font("Author", "DejaVu Sans", 10, False, slate)
    for name, size in (
        ("Heading 1", 18.0),
        ("Heading 2", 14.0),
        ("Heading 3", 12.0),
        ("Heading 4", 11.0),
    ):
        if name in styles:
            st = set_font(name, "DejaVu Sans", size, True, navy if name != "Heading 4" else slate)
            st.paragraph_format.keep_with_next = True
            st.paragraph_format.space_before = Pt(11 if name == "Heading 1" else 7)
            st.paragraph_format.space_after = Pt(3.5)
    if "Caption" in styles:
        cap = set_font("Caption", "DejaVu Sans", 8.4, False, slate)
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(7)
    for name in ("Block Text", "Quote"):
        if name in styles:
            st = set_font(name, "DejaVu Serif", 11.0, False, slate)
            st.paragraph_format.left_indent = Inches(0.28)
            st.paragraph_format.right_indent = Inches(0.18)
            st.paragraph_format.space_before = Pt(4)
            st.paragraph_format.space_after = Pt(5)

    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            p.paragraph_format.keep_with_next = True
        if any(
            p.text.startswith(prefix)
            for prefix in (
                "Candidate integration manuscript",
                "Repository basis.",
                "Scientific status.",
                "Licensing intent.",
            )
        ):
            for r in p.runs:
                r.font.name = "DejaVu Sans"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "DejaVu Sans")
                r.font.size = Pt(8.8)
                r.font.color.rgb = slate
        for r in p.runs:
            if r.font.name is None:
                r.font.name = "DejaVu Serif"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "DejaVu Serif")
        p_pr = p._p.get_or_add_pPr()
        if p_pr.find(qn("w:widowControl")) is None:
            p_pr.append(OxmlElement("w:widowControl"))

    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.clear()
        hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.7))
        r = hp.add_run("SPPT / ASTRA")
        r.font.name = "DejaVu Sans"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "DejaVu Sans")
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = navy
        hp.add_run("\t")
        r = hp.add_run("v1.0.8 candidate")
        r.font.name = "DejaVu Sans"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "DejaVu Sans")
        r.font.size = Pt(8)
        r.font.color.rgb = slate
        p_pr = hp._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for key, val in (
            ("w:val", "single"),
            ("w:sz", "4"),
            ("w:space", "2"),
            ("w:color", "B99822"),
        ):
            bottom.set(qn(key), val)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        fp = sec.footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        fp._p.append(fld)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        try:
            table.style = "Light Shading Accent 1"
        except Exception:
            pass
        for ri, row in enumerate(table.rows):
            if ri == 0:
                tr_pr = row._tr.get_or_add_trPr()
                marker = OxmlElement("w:tblHeader")
                marker.set(qn("w:val"), "true")
                tr_pr.append(marker)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(1.5)
                    p.paragraph_format.line_spacing = 1.0
                    for r in p.runs:
                        r.font.name = "DejaVu Sans"
                        r._element.rPr.rFonts.set(qn("w:eastAsia"), "DejaVu Sans")
                        r.font.size = Pt(8.2 if len(table.columns) >= 5 else 8.8)
                        if ri == 0:
                            r.font.bold = True
                            r.font.color.rgb = navy

    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1":
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            for key, val in (
                ("w:val", "single"),
                ("w:sz", "6"),
                ("w:space", "3"),
                ("w:color", "B99822"),
            ):
                bottom.set(qn(key), val)
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
        if p._p.xpath(".//w:drawing"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_after = Pt(1)

    # Contain every placed figure within the text block while preserving
    # source aspect ratio. Fixed-height boxes are forbidden because they clip
    # labels when text expands.
    max_w = int(Inches(6.45))
    max_h = int(Inches(6.65))
    for shape in doc.inline_shapes:
        old_w = int(shape.width)
        old_h = int(shape.height)
        if old_w <= 0 or old_h <= 0:
            continue
        scale = min(max_w / old_w, max_h / old_h)
        shape.width = int(old_w * scale)
        shape.height = int(old_h * scale)

    doc.save(str(path_out))


def prepend_text_cover(path_in: Path, path_out: Path) -> None:
    """Prepend a plain, selectable cover page to the styled DOCX."""
    doc = Document(str(path_in))
    body = doc._element.body
    cover_elements = []

    def add(
        text: str = "",
        *,
        size: float = 11,
        bold: bool = False,
        color: RGBColor | None = None,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before: float = 0,
        after: float = 0,
    ):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.font.name = "DejaVu Sans"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "DejaVu Sans")
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        cover_elements.append(p._p)
        return p

    navy = RGBColor(0x07, 0x1B, 0x2A)
    gold = RGBColor(0xB9, 0x98, 0x22)
    slate = RGBColor(0x3B, 0x4B, 0x56)
    add("SPPT / ASTRA", size=15, bold=True, color=gold, before=62, after=36)
    add("v1.0.8", size=36, bold=True, color=navy, after=3)
    add("ENDOGENOUS VISIBILITY", size=20, bold=True, color=navy, after=16)
    add("--------------------------------", size=11, color=gold, after=22)
    add("Source-coupled transducers", size=20, bold=True, color=navy, after=0)
    add("and transformed cosmic archives", size=20, bold=True, color=navy, after=22)
    add("Self-detuning media · cross-channel rescue", size=10.5, color=slate, after=2)
    add("Catastrophic archives · global certificates", size=10.5, color=slate, after=54)
    add("Candidate successor draft", size=9.5, color=slate, after=1)
    add("Not peer reviewed · no empirical planetary validation", size=9.5, color=slate, after=1)
    add(
        "Repository-visible; no tag, release, Pages, DOI, or Zenodo action",
        size=9.5,
        color=slate,
        after=24,
    )
    add("16 August 2026     ·     AD ASTRA PER ASPERA", size=9.5, bold=True, color=gold)
    pbreak = doc.add_paragraph()
    pbreak.add_run().add_break(WD_BREAK.PAGE)
    cover_elements.append(pbreak._p)

    for el in cover_elements:
        body.remove(el)
    for el in reversed(cover_elements):
        body.insert(0, el)
    doc.save(str(path_out))


def build_docx() -> None:
    raw = WORK / "candidate_unstyled.docx"
    styled = WORK / "candidate_styled.docx"
    moved = WORK / "candidate_moved.docx"
    body_md = WORK / "body_without_cover_docx.md"
    body_md.write_text(prepare_docx_body_text(MD.read_text(encoding="utf-8")), encoding="utf-8")
    run_pandoc(
        str(body_md),
        "--from",
        "markdown+raw_tex",
        "--citeproc",
        "--bibliography",
        str(SOURCE / "references.bib"),
        f"--resource-path={os.pathsep.join((str(SOURCE), str(FIGURES)))}",
        "-o",
        str(raw),
    )
    style_docx(raw, styled)
    prepend_text_cover(styled, moved)
    out = ROOT / "ASTRA_SPPT_v1.0.8_Endogenous_Visibility_Candidate.docx"
    # Safe accessibility fixes, then privacy metadata scrub.
    tools_directory = Path(os.environ.get("ASTRA_DOCX_TOOLS_DIR", "")).resolve()
    a11y_script = tools_directory / "a11y_audit.py"
    privacy_script = tools_directory / "privacy_scrub.py"
    if not a11y_script.is_file() or not privacy_script.is_file():
        raise RuntimeError("ASTRA_DOCX_TOOLS_DIR must contain a11y_audit.py and privacy_scrub.py")
    a11y = WORK / "candidate_a11y.docx"
    run(
        sys.executable,
        str(a11y_script),
        str(moved),
        "--fix_table_headers",
        "first_row",
        "--fix_image_alt",
        "from_filename",
        "--out",
        str(a11y),
        "--out_json",
        str(WORK / "docx_a11y_pre.json"),
    )
    run(
        sys.executable,
        str(privacy_script),
        str(a11y),
        "--out",
        str(out),
    )
    normalize_docx_archive(out)
    report_path = VERIFY / "docx_a11y_report.json"
    run(
        sys.executable,
        str(a11y_script),
        str(out),
        "--out_json",
        str(report_path),
    )
    report = json.loads(report_path.read_text(encoding="utf-8"))
    report["file"] = out.name
    report_path.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def build_peer_review_pdf() -> None:
    """Build a journal-style review copy: 12 pt, double spaced, one column, continuous line numbers."""
    review_md = WORK / "peer_review_source.md"
    text = MD.read_text(encoding="utf-8")
    text = prepare_body_text(text)
    text = text.replace("fontsize: 11pt", "fontsize: 12pt")
    text = text.replace("geometry: margin=0.78in", "geometry: margin=1in")
    text = text.replace(r"\setstretch{1.18}", r"\doublespacing")
    text = text.replace(
        r"\usepackage{setspace}",
        "\\usepackage{setspace}\n    \\usepackage[switch]{lineno}\n    \\linenumbers",
    )
    text = text.replace(r"\pagestyle{fancy}", r"\pagestyle{plain}")
    review_md.write_text(text, encoding="utf-8")
    raw = WORK / "ASTRA_SPPT_v1.0.8_Endogenous_Visibility_Candidate_Peer_Review.pdf"
    run_pandoc(
        str(review_md),
        "--from",
        "markdown+raw_tex",
        "--citeproc",
        "--bibliography",
        str(SOURCE / "references.bib"),
        f"--pdf-engine={tex_engine()}",
        f"--resource-path={os.pathsep.join((str(SOURCE), str(FIGURES)))}",
        "-o",
        str(raw),
    )
    _sanitize_pdf(
        raw,
        ROOT / "ASTRA_SPPT_v1.0.8_Endogenous_Visibility_Candidate_Peer_Review.pdf",
        "SPPT / ASTRA v1.0.8 Candidate: Peer-Review Edition",
    )


def _embedded_font_css() -> str:
    families = {
        "ASTRA Serif": font_manager.FontProperties(family="DejaVu Serif"),
        "ASTRA Sans": font_manager.FontProperties(family="DejaVu Sans"),
        "ASTRA Mono": font_manager.FontProperties(family="DejaVu Sans Mono"),
    }
    rules: list[str] = []
    for family, properties in families.items():
        path = Path(font_manager.findfont(properties, fallback_to_default=False))
        payload = base64.b64encode(path.read_bytes()).decode("ascii")
        rules.append(
            "@font-face {"
            f"font-family: '{family}'; font-style: normal; font-weight: 400;"
            f"src: url(data:font/ttf;base64,{payload}) format('truetype');"
            "}"
        )
    return "\n".join(rules)


def _candidate_css() -> str:
    return (
        "<style>"
        + _embedded_font_css()
        + """
        @page { size: Letter; margin: 0.72in 0.78in 0.70in; }
        :root { color-scheme: light; }
        html { background: #ffffff; }
        body {
          color: #11212b;
          font-family: 'ASTRA Serif', serif;
          font-size: 11pt;
          line-height: 1.28;
          margin: 0 auto;
          max-width: 7in;
        }
        h1, h2, h3, h4, nav, figcaption, table {
          font-family: 'ASTRA Sans', sans-serif;
        }
        h1, h2, h3, h4 { color: #071b2a; break-after: avoid-page; }
        h1 { font-size: 20pt; margin-top: 1.1em; }
        h2 { font-size: 15pt; }
        h3 { font-size: 12.5pt; }
        p, li { orphans: 3; widows: 3; }
        a { color: #1d5f96; text-decoration-thickness: 0.08em; }
        img, svg { display: block; height: auto; margin: 0.6em auto; max-width: 100%; }
        figure { break-inside: avoid-page; margin: 1em 0; }
        figcaption { font-size: 9pt; line-height: 1.3; }
        table { border-collapse: collapse; font-size: 8.5pt; width: 100%; }
        th, td { border: 0.5pt solid #61717c; padding: 0.28em 0.35em; vertical-align: top; }
        th { background: #e8edf0; color: #071b2a; text-align: left; }
        pre, code { font-family: 'ASTRA Mono', monospace; overflow-wrap: anywhere; }
        math[display='block'] { display: block; margin: 0.8em auto; max-width: 100%; }
        nav { border: 1pt solid #aeb9bf; padding: 0.7em 1em; }
        @media print {
          body { max-width: none; }
          nav { break-after: page; }
          a { color: #071b2a; }
        }
        """
        + "</style>"
    )


def _sanitize_pdf(source: Path, destination: Path, title: str) -> None:
    with pikepdf.open(source) as pdf:
        for key in list(pdf.docinfo):
            del pdf.docinfo[key]
        pdf.docinfo.update(
            {
                "/Title": title,
                "/Author": "Jacko T.",
                "/Subject": "Unpromoted SPPT / ASTRA v1.0.8 candidate",
                "/Creator": "ASTRA v1.0.8 candidate document builder",
                "/Producer": "pikepdf 10.11.0",
                "/CreationDate": PDF_DATE,
                "/ModDate": PDF_DATE,
            }
        )
        pdf.Root["/Lang"] = pikepdf.String("en-US")
        pdf.Root["/ViewerPreferences"] = pikepdf.Dictionary(DisplayDocTitle=True)
        if "/Metadata" in pdf.Root:
            del pdf.Root["/Metadata"]
        with pdf.open_metadata(set_pikepdf_as_editor=False, update_docinfo=False) as metadata:
            metadata["dc:title"] = title
            metadata["dc:creator"] = ["Jacko T."]
            metadata["dc:language"] = ["en-US"]
            metadata["xmp:CreateDate"] = BUILD_EPOCH
            metadata["xmp:ModifyDate"] = BUILD_EPOCH
            metadata["xmp:MetadataDate"] = BUILD_EPOCH
        if "/ID" in pdf.trailer:
            del pdf.trailer["/ID"]
        pdf.save(
            destination,
            deterministic_id=True,
            object_stream_mode=pikepdf.ObjectStreamMode.generate,
            compress_streams=True,
        )


def _build_html_and_tagged_pdf(
    source: Path,
    html_output: Path,
    pdf_output: Path,
    title: str,
    *,
    bibliography: bool,
) -> None:
    temporary_html = WORK / f"{source.stem}.html"
    arguments = [
        str(source),
        "--from",
        "markdown+raw_tex",
        "--to",
        "html5",
        "--standalone",
        "--toc",
        "--toc-depth=3",
        "--mathml",
        "--embed-resources",
        f"--resource-path={os.pathsep.join((str(SOURCE), str(FIGURES), str(ROOT)))}",
        f"--metadata=pagetitle:{title}",
        "--metadata=lang:en-US",
    ]
    if bibliography:
        arguments.extend(
            [
                "--citeproc",
                "--bibliography",
                str(SOURCE / "references.bib"),
            ]
        )
    arguments.extend(["-o", str(temporary_html)])
    run_pandoc(*arguments)
    html = temporary_html.read_text(encoding="utf-8")
    if "</head>" not in html:
        raise RuntimeError(f"Pandoc output lacks a head element: {source}")
    html = html.replace("</head>", _candidate_css() + "</head>", 1)
    html_output.write_text(html, encoding="utf-8", newline="\n")

    raw_pdf = WORK / f"{source.stem}-raw.pdf"
    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=True, args=["--disable-gpu"])
        page = browser.new_page()
        page.goto(html_output.resolve().as_uri(), wait_until="networkidle")
        page.emulate_media(media="print")
        page.evaluate("document.fonts.ready")
        page.pdf(
            path=str(raw_pdf),
            format="Letter",
            print_background=True,
            prefer_css_page_size=True,
            tagged=True,
            outline=True,
        )
        browser.close()
    _sanitize_pdf(raw_pdf, pdf_output, title)


def build_tagged_reading_edition() -> None:
    _build_html_and_tagged_pdf(
        MD,
        WORK / "ASTRA_SPPT_v1.0.8_Endogenous_Visibility_Candidate.html",
        ROOT / "ASTRA_SPPT_v1.0.8_Endogenous_Visibility_Candidate_Tagged_Reading_Edition.pdf",
        "SPPT / ASTRA v1.0.8 Candidate: Endogenous Visibility",
        bibliography=True,
    )


def build_verification_report_pdf() -> None:
    _build_html_and_tagged_pdf(
        VERIFY / "verification_report.md",
        WORK / "verification_report.html",
        VERIFY / "ASTRA_SPPT_v1.0.8_Candidate_Verification_Report.pdf",
        "SPPT / ASTRA v1.0.8 Candidate Verification Report",
        bibliography=False,
    )


def main() -> None:
    global WORK

    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--only",
        choices=("all", "verification-report"),
        default="all",
        help="build all editions, or rebuild only the verification-report PDF",
    )
    arguments = parser.parse_args()

    VERIFY.mkdir(exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="astra-v108-build-") as directory:
        WORK = Path(directory)
        if arguments.only == "all":
            build_pdf()
            build_peer_review_pdf()
            build_docx()
            build_tagged_reading_edition()
        build_verification_report_pdf()

    print(f"generated candidate editions for fixed build epoch {BUILD_EPOCH}")


if __name__ == "__main__":
    main()
